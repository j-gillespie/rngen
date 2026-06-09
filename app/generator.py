"""Build Markdown release notes from categorized content and template metadata."""

from __future__ import annotations

from pathlib import Path

from docx import Document

from app.categorizer import CategorizedNotes, categorize_details

TEMPLATE_PATH = Path(__file__).resolve().parent.parent / "templates" / "Release Notes.docx"
NONE_TEXT = "None for this release."
DEFAULT_SUPPORT = (
    "For assistance with this product, contact the ACME Technical Support team "
    "through your standard service desk channel or email support@acme.example.com. "
    "Include the product name, version, and a detailed description of the issue."
)


def _load_template_sections() -> list[str]:
    document = Document(TEMPLATE_PATH)
    sections: list[str] = []
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            sections.append(text)
    return sections


def _section_items(items: list[str], bullet: bool = False) -> str:
    if not items:
        return NONE_TEXT
    if bullet:
        return "\n".join(f"- {item}" for item in items)
    if len(items) == 1:
        return items[0]
    return "\n".join(f"- {item}" for item in items)


def _format_date(iso_date: str) -> str:
    year, month, day = iso_date.split("-")
    return f"{month}/{day}/{year}"


def generate_markdown(
    product: str,
    version: str,
    release_date: str,
    details: str,
) -> str:
    notes = categorize_details(details)
    formatted_date = _format_date(release_date)

    sections = [
        (f"# {product} {version} Release Notes", f"**Release Date:** {formatted_date}"),
        ("## Overview", _section_items(notes.overview)),
        ("## New Features", _section_items(notes.new_features, bullet=True)),
        ("## Resolved Issues", _section_items(notes.resolved_issues, bullet=True)),
        ("## Known Issues", _section_items(notes.known_issues, bullet=True)),
        (
            "## System Requirements & Compatibility",
            _section_items(notes.system_requirements, bullet=True),
        ),
        ("## Installation", _section_items(notes.installation, bullet=True)),
        (
            "## Technical Support",
            _section_items(notes.technical_support)
            if notes.technical_support
            else DEFAULT_SUPPORT,
        ),
    ]

    lines: list[str] = []
    for index, (heading, content) in enumerate(sections):
        lines.extend([heading, content])
        if index < len(sections) - 1:
            lines.append("")

    return "\n".join(lines)


def populate_docx_template(
    product: str,
    version: str,
    release_date: str,
    details: str,
    output_path: Path,
) -> Path:
    """Optional: write a populated .docx using the Word template placeholders."""
    notes = categorize_details(details)
    formatted_date = _format_date(release_date)
    document = Document(TEMPLATE_PATH)

    replacements = {
        "[PRODUCT NAME]": product,
        "[VERSION]": version,
        "[DATE]": formatted_date,
    }

    section_content = {
        "Overview": _section_items(notes.overview),
        "New Features": _section_items(notes.new_features, bullet=True),
        "Resolved Issues": _section_items(notes.resolved_issues, bullet=True),
        "Known Issues": _section_items(notes.known_issues, bullet=True),
        "System Requirements & Compatibility": _section_items(
            notes.system_requirements, bullet=True
        ),
        "Installation": _section_items(notes.installation, bullet=True),
        "Technical Support": (
            _section_items(notes.technical_support)
            if notes.technical_support
            else DEFAULT_SUPPORT
        ),
    }

    current_section: str | None = None
    for paragraph in document.paragraphs:
        for placeholder, value in replacements.items():
            if placeholder in paragraph.text:
                paragraph.text = paragraph.text.replace(placeholder, value)

        heading = paragraph.text.strip()
        if heading in section_content:
            current_section = heading
            continue

        if current_section and heading and heading not in section_content:
            if paragraph.text.strip() and not paragraph.text.startswith("["):
                paragraph.text = section_content[current_section]
                current_section = None

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)
    return output_path
